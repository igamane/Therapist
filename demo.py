import os
import openai
import time
import streamlit as st
from dotenv import load_dotenv
import time
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from googleapiclient.discovery import build
from google.oauth2 import service_account
import json
import sys


# Load environment variables from .env file
load_dotenv()

# Set OpenAI API key
openai.api_key = os.getenv("OPENAI_API_KEY")
document_file = os.getenv("DOCUMENT_FILE")
FOLDER_ID = os.getenv("FOLDER_ID")
drive_file_name = os.getenv("DRIVE_FILE_NAME")

# Load the Word document using the absolute path
doc = Document(document_file)

# Create a client instance
client = openai.Client()

SCOPES = ['https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = 'client_secrets.json'

def authenticate():
    creds = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    return creds

def find_file_id(service, file_name):
    results = service.files().list(
        q=f"name='{file_name}' and parents in '{FOLDER_ID}' and trashed=false",
        fields="files(id)"
    ).execute()
    items = results.get('files', [])
    if items:
        return items[0]['id']
    else:
        return None

def delete_file(service, file_id):
    service.files().delete(fileId=file_id).execute()

def upload_file(file_path):
    creds = authenticate()
    service = build('drive', 'v3', credentials=creds)

    file_name = drive_file_name  # Replace with the desired file name
    existing_file_id = find_file_id(service, file_name)

    if existing_file_id:
        delete_file(service, existing_file_id)
        print(f"Existing file '{file_name}' deleted.")

    file_metadata = {
        'name': file_name,
        'parents': [FOLDER_ID]
    }

    file = service.files().create(
        body=file_metadata,
        media_body=file_path
    ).execute()

    print(f"File '{file_name}' uploaded successfully.")

# %%%%%%%%%%%%%%%%%% Document Updating Functions %%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
def add_paragraph_after_header(header_text, new_paragraph):
    found_header = False
    for idx, paragraph in enumerate(doc.paragraphs):
        if header_text in paragraph.text:
            found_header = True
        elif found_header:
             # Insert the new paragraph after the found header
            new_para = doc.paragraphs[idx].insert_paragraph_before(new_paragraph)
            
            # Highlight the added paragraph in yellow
            for run in new_para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # Add a paragraph with no text after the table to create space
            doc.add_paragraph()
            doc.save(document_file)
            upload_file(document_file)
            return f"Document update: the paragraph has been added after the header'{header_text}' - tell the user what you have added"

    if not found_header:
        return f"Header '{header_text}' not found in the document."
  
def add_row_to_table_by_index(table_index, row_data):
    total_tables = len(doc.tables)

    if table_index < 0 or table_index >= total_tables:
        return f"Table index '{table_index}' is out of range."

    table = doc.tables[table_index]

    # Add a row to the table with the specified data
    new_row = table.add_row().cells
    try:
        for i, cell_data in enumerate(row_data):
            new_row[i].text = str(cell_data)

            # Highlight the added cell content in yellow
            for paragraph in new_row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except IndexError:
        return "The column count is out of the table range."
    # Handle the out of range column scenario as per your requirement

    doc.save(document_file)
    upload_file(document_file)
    return "Document update: row has been added to the table - draw the row you have added to the user"

def create_new_table(data):
    num_rows = len(data)
    if num_rows == 0:
        return "No data provided."

    num_cols = len(data[0])

    # Add a table to the document based on the size of the data
    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Retrieve the style of the previous table if it exists
    if len(doc.tables) > 1:  # Assuming the previous table is at index 0
        previous_table_style = doc.tables[0].style
        table.style = previous_table_style.name  # Apply the style to the new table

    # Insert values into the cells of the table
    for i in range(num_rows):
        for j in range(num_cols):
            cell = table.cell(i, j)
            cell.text = str(data[i][j])  # Set the text for each cell
            
            # Highlight the added cell content in yellow
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            if i == 0:  # Bold the text in the first row
                cell.paragraphs[0].runs[0].bold = True

    # Add a paragraph with no text after the table to create space
    doc.add_paragraph()
    doc.save(document_file)
    upload_file(document_file)

    return "Document update: table has been created - tell the user the table content"


def add_new_section(header_text, section_content):
    # Check if 'Heading 1' style exists in the document
    styles = doc.styles
    heading_style_exists = any(style.name == 'Heading 1' and style.type == WD_STYLE_TYPE.PARAGRAPH for style in styles)

    # If 'Heading 1' style doesn't exist, create it
    if not heading_style_exists:
        style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.bold = True
        font.size = Pt(16)

    # Add a heading with the specified text using the 'Heading 1' style
    doc.add_heading(header_text, level=1)

    # Check if '**' exists in the section_content
    if '**' not in section_content:
        # If '**' is not present, add the entire section as regular text
        p = doc.add_paragraph(section_content)
        # Highlight entire section in yellow
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        # Split the section content by '**'
        sections = section_content.split('**')

        # Add content for the new section
        for i, text in enumerate(sections):
            p = doc.add_paragraph()
            p.add_run(text)

            # Highlight the added section in yellow
            for run in p.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Add a paragraph with no text after the table to create space
    doc.add_paragraph()
    doc.save(document_file)
    upload_file(document_file)

    return "Document update: new section has been added - tell the user what you have added"

# Load assistant details from JSON file
try:
    with open('starter_questions.json', 'r') as json_file:
        assistant_questions = json.load(json_file)
except FileNotFoundError:
    print("Error: starter_questions.json file not found.")
    sys.exit(1)
except json.JSONDecodeError:
    print("Error: starter_questions.json is not a valid JSON file.")
    sys.exit(1)

# Function to update starter questions based on the assistant name
def update_starter_questions(assistant_id):
    for assistant, details in assistant_questions.items():
        if details.get("id") == assistant_id:
            return details.get("starter_questions", [])

def get_response(assistant_id):
    # Check if 'messages' key is not in session_state
    if "messages" not in st.session_state:
    # If not present, initialize 'messages' as an empty list
        st.session_state.messages = []
    # Iterate through messages in session_state
    for message in st.session_state.messages:
    # Display message content in the chat UI based on the role
        with st.chat_message(message["role"]):
            st.markdown(message["content"])    
    if not st.session_state.get("starter_displayed", False):
    # starter questions
        starter_questions = update_starter_questions(assistant_id)
        
        placeholder = st.empty()

        col1, col2 = placeholder.columns(2)

        clicked_question = False

        question_v = ""
        with col1:
            for idx, question in enumerate(starter_questions[:2]):
                button_key = f"btn_col1_{idx + 1}"  # Unique key for column 1 buttons
                if st.button(question, key=button_key):
                    question_v = question
                    clicked_question = True
                    # Replace user prompt with the starter question when clicked
                    break  # Exit the loop if a question is clicked

        with col2:
            for idx, question in enumerate(starter_questions[2:]):
                button_key = f"btn_col2_{idx + 1}"  # Unique key for column 2 buttons
                if st.button(question, key=button_key):
                    question_v = question
                    clicked_question = True
                    # Replace user prompt with the starter question when clicked
                    break  # Exit the loop if a question is clicked

        if clicked_question:
            placeholder.empty()
            st.session_state.messages.append({"role": "user", "content": question_v})
            with st.chat_message("user"):
                st.markdown(question_v)
            # Process the assistant's response using the starter question
            st.session_state.starter_displayed = True
            with st.spinner("Thinking..."):
                getResponse(assistant_id, question_v)
    # Get user input from chat and proceed if a prompt is entered
    if prompt := st.chat_input("Enter your message here"):
        if not st.session_state.get("starter_displayed", False):
            placeholder.empty()
            st.session_state.starter_displayed = True
        # Add user input as a message to session_state
        st.session_state.messages.append({"role": "user", "content": prompt})
        # Display user's message in the chat UI
        with st.chat_message("user"):
            st.markdown(prompt)
        # Process the assistant's response
        with st.spinner("Thinking..."):
            getResponse(assistant_id, prompt)


def getResponse(assistant_id, prompt):
    if "thread_id" not in st.session_state:
        thread = client.beta.threads.create()
        st.session_state.thread_id = thread.id
    thread_id = st.session_state.thread_id    


    for files in st.session_state.uploaded_files_list:
        client.beta.assistants.files.create(
            assistant_id=assistant_id,
            file_id=files
        )

    message = client.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content= prompt,
    )

    run = client.beta.threads.runs.create(
        thread_id=thread_id,
        assistant_id=assistant_id,
        )
    
    # Define a dispatch table
    function_dispatch_table = {
        "add_paragraph_after_header": add_paragraph_after_header,
        "add_row_to_table_by_index": add_row_to_table_by_index,
        "create_new_table": create_new_table,
        "add_new_section": add_new_section
    }

    while True:
        # Wait for 5 seconds
        time.sleep(0.5)

        # Retrieve the run status
        run_status = client.beta.threads.runs.retrieve(
            thread_id=thread_id,
            run_id=run.id
        )
        print(run_status.status)
        # If run is completed, get messages
        if run_status.status == 'completed':
            break
        elif run_status.status == 'failed':
            print(run_status)
            break
        elif run_status.status == 'requires_action':
            required_actions = run_status.required_action.submit_tool_outputs.model_dump()
            tool_outputs = []
            import json
            for action in required_actions["tool_calls"]:
                func_name = action['function']['name']
                arguments = json.loads(action['function']['arguments'])
                print(func_name)
                print(arguments)

                func = function_dispatch_table.get(func_name)

                if func:
                    result = func(**arguments)
                    # Ensure the output is a JSON string
                    output = json.dumps(result) if not isinstance(result, str) else result
                    print(output)
                    tool_outputs.append({
                        "tool_call_id": action["id"],
                        "output": output
                    })
                else:
                    print(f"Function {func_name} not found")
            
            client.beta.threads.runs.submit_tool_outputs(
                thread_id=thread_id,
                run_id=run.id,
                tool_outputs=tool_outputs
            )
        else:
            time.sleep(0.5)

    messages = client.beta.threads.messages.list(
        thread_id=thread_id
    )

    last_message = messages.data[0].content[0].text.value

    st.session_state.messages.append({"role": "assistant", "content": last_message})
    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        message_placeholder.markdown(last_message)

with open('assistants.json') as json_file:
    assistant_id_map = json.load(json_file)

def get_assistant_id(assistant_name):
    return assistant_id_map.get(assistant_name)

def send_to_openai(file):
    try:
        local_file_path = os.path.join(os.getcwd(), file.name)  # Save the file in the working directory
        with open(local_file_path, "wb") as f:
            f.write(file.read())
        
        # Send file path to OpenAI
        with open(local_file_path, "rb") as file:
            response = client.files.create(
                file=file,
                purpose="assistants"
            )
                    
        os.remove(local_file_path)

        return response
    except Exception as e:
        print(f"Error: {str(e)}")

def parentPage():
    st.title('Parent AI assistant')

    st.sidebar.title("Document Processing")
    assistant_name = "Parent"
    uploaded_files = st.sidebar.file_uploader("Upload files", accept_multiple_files=True)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_identifier = uploaded_file.name
            if file_identifier not in st.session_state.sent_files:
                retrieval_file = send_to_openai(uploaded_file)
                st.session_state.uploaded_files_list.append(retrieval_file.id)  # Store uploaded file in session_state
                st.session_state.sent_files.add(file_identifier)

    if assistant_name != "Select Category":
        # Reset the conversation & starter questions, if the assistant has been changed
        if st.session_state.get("current_assistant") != assistant_name:
            st.session_state.starter_displayed = False
            st.session_state.messages = []
            st.session_state.current_assistant = assistant_name
        assistant_id = get_assistant_id(assistant_name)
        print(assistant_id)
        get_response(assistant_id)

    # Add download button for myfile.csv
    with open(document_file, 'rb') as f:
        st.sidebar.download_button('Download The Updated Achilles Guide', f, file_name=document_file, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
   
    if st.sidebar.button("Return to Main Page"):
        st.session_state.runpage = main_page
        st.rerun()

def therapistPage():
    st.title('Therapist AI assistant')

    st.sidebar.title("Document Processing")
    assistant_name = "Therapist"
    uploaded_files = st.sidebar.file_uploader("Upload files", accept_multiple_files=True)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_identifier = uploaded_file.name
            if file_identifier not in st.session_state.sent_files:
                retrieval_file = send_to_openai(uploaded_file)
                st.session_state.uploaded_files_list.append(retrieval_file.id)  # Store uploaded file in session_state
                st.session_state.sent_files.add(file_identifier)

    if assistant_name != "Select Category":
        # Reset the conversation & starter questions, if the assistant has been changed
        if st.session_state.get("current_assistant") != assistant_name:
            st.session_state.starter_displayed = False
            st.session_state.messages = []
            st.session_state.current_assistant = assistant_name
        assistant_id = get_assistant_id(assistant_name)
        print(assistant_id)
        get_response(assistant_id)

    # Add download button for myfile.csv
    with open(document_file, 'rb') as f:
        st.sidebar.download_button('Download The Updated Achilles Guide', f, file_name=document_file, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
   
    if st.sidebar.button("Return to Main Page"):
        st.session_state.runpage = main_page
        st.rerun()

def teacherPage():
    st.title('Teacher AI assistant')

    st.sidebar.title("Document Processing")
    assistant_name = "Teacher"
    uploaded_files = st.sidebar.file_uploader("Upload files", accept_multiple_files=True)

    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_identifier = uploaded_file.name
            if file_identifier not in st.session_state.sent_files:
                retrieval_file = send_to_openai(uploaded_file)
                st.session_state.uploaded_files_list.append(retrieval_file.id)  # Store uploaded file in session_state
                st.session_state.sent_files.add(file_identifier)

    if assistant_name != "Select Category":
        # Reset the conversation & starter questions, if the assistant has been changed
        if st.session_state.get("current_assistant") != assistant_name:
            st.session_state.starter_displayed = False
            st.session_state.messages = []
            st.session_state.current_assistant = assistant_name
        assistant_id = get_assistant_id(assistant_name)
        print(assistant_id)
        get_response(assistant_id)

    # Add download button for myfile.csv
    with open(document_file, 'rb') as f:
        st.sidebar.download_button('Download The Updated Achilles Guide', f, file_name=document_file, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
   
    if st.sidebar.button("Return to Main Page"):
        st.session_state.runpage = main_page
        st.rerun()

def main_page():
    st.title("Select a profile")
    btn1 = st.button("Parent")
    btn2 = st.button("Therapist")
    btn3 = st.button("Teacher")

    if btn1:
        st.session_state.runpage = parentPage
        st.rerun()

    if btn2:
        st.session_state.runpage = therapistPage
        st.rerun()
    
    if btn3:
        st.session_state.runpage = teacherPage
        st.rerun()

def main():
    if 'runpage' not in st.session_state:
        st.session_state.runpage = main_page

    # Initialize uploaded_files_list in st.session_state if it doesn't exist
    if "uploaded_files_list" not in st.session_state:
        st.session_state.uploaded_files_list = []
        st.session_state.sent_files = set()

    st.session_state.runpage()

if __name__ == "__main__":
    main()
